from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def create_advanced_word_document():
    """创建更复杂的Word自动编号测试文档"""
    doc = Document()
    
    # 设置文档样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    
    # 添加主标题
    doc.add_heading('Word自动编号识别测试文档（高级版）', 0)
    
    # 添加说明
    doc.add_paragraph('这个文档包含Word的各种自动编号格式，用于测试识别算法的准确性：')
    
    # 第一部分：多级列表
    doc.add_heading('第一部分：多级列表', level=1)
    
    # 创建多级列表
    # 一级列表
    p1 = doc.add_paragraph('第一章 绪论')
    p1.style = 'List Number'
    
    p2 = doc.add_paragraph('1.1 研究背景')
    p2.style = 'List Number 2'
    
    p3 = doc.add_paragraph('1.1.1 国内研究现状')
    p3.style = 'List Number 3'
    
    p4 = doc.add_paragraph('1.1.2 国外研究现状')
    p4.style = 'List Number 3'
    
    p5 = doc.add_paragraph('1.2 研究意义')
    p5.style = 'List Number 2'
    
    # 继续一级列表
    p6 = doc.add_paragraph('第二章 文献综述')
    p6.style = 'List Number'
    
    p7 = doc.add_paragraph('2.1 理论基础')
    p7.style = 'List Number 2'
    
    p8 = doc.add_paragraph('2.2 研究方法')
    p8.style = 'List Number 2'
    
    # 第二部分：字母编号
    doc.add_heading('第二部分：字母编号', level=1)
    
    # 大写字母编号
    p9 = doc.add_paragraph('A. 项目概述')
    p9.style = 'List Number'
    
    p10 = doc.add_paragraph('B. 需求分析')
    p10.style = 'List Number'
    
    p11 = doc.add_paragraph('C. 系统设计')
    p11.style = 'List Number'
    
    # 小写字母编号
    p12 = doc.add_paragraph('a) 功能模块')
    p12.style = 'List Number 2'
    
    p13 = doc.add_paragraph('b) 接口设计')
    p13.style = 'List Number 2'
    
    # 第三部分：罗马数字
    doc.add_heading('第三部分：罗马数字编号', level=1)
    
    p14 = doc.add_paragraph('I. 第一阶段')
    p14.style = 'List Number'
    
    p15 = doc.add_paragraph('II. 第二阶段')
    p15.style = 'List Number'
    
    p16 = doc.add_paragraph('III. 第三阶段')
    p16.style = 'List Number'
    
    # 第四部分：中文数字编号
    doc.add_heading('第四部分：中文数字编号', level=1)
    
    p17 = doc.add_paragraph('第一条 总则')
    p17.style = 'List Number'
    
    p18 = doc.add_paragraph('第二条 细则')
    p18.style = 'List Number'
    
    p19 = doc.add_paragraph('第三条 附则')
    p19.style = 'List Number'
    
    # 第五部分：混合编号
    doc.add_heading('第五部分：混合编号测试', level=1)
    
    # 创建复杂的混合编号结构
    p20 = doc.add_paragraph('1. 主要功能')
    p20.style = 'List Number'
    
    p21 = doc.add_paragraph('1.1 用户管理')
    p21.style = 'List Number 2'
    
    p22 = doc.add_paragraph('1.1.1 注册功能')
    p22.style = 'List Number 3'
    
    p23 = doc.add_paragraph('1.1.2 登录功能')
    p23.style = 'List Number 3'
    
    p24 = doc.add_paragraph('1.2 权限管理')
    p24.style = 'List Number 2'
    
    p25 = doc.add_paragraph('2. 系统设置')
    p25.style = 'List Number'
    
    p26 = doc.add_paragraph('2.1 基础配置')
    p26.style = 'List Number 2'
    
    p27 = doc.add_paragraph('2.2 高级选项')
    p27.style = 'List Number 2'
    
    # 第六部分：特殊格式
    doc.add_heading('第六部分：特殊格式测试', level=1)
    
    # 括号编号
    p28 = doc.add_paragraph('(1) 括号格式一')
    p28.style = 'List Number'
    
    p29 = doc.add_paragraph('(2) 括号格式二')
    p29.style = 'List Number'
    
    # 点号编号
    p30 = doc.add_paragraph('1. 点号格式')
    p30.style = 'List Number'
    
    p31 = doc.add_paragraph('2. 点号格式')
    p31.style = 'List Number'
    
    # 第七部分：实际应用场景
    doc.add_heading('第七部分：实际应用场景', level=1)
    
    # 模拟合同条款
    p32 = doc.add_paragraph('第一条 合同标的')
    p32.style = 'List Number'
    
    p33 = doc.add_paragraph('1.1 产品规格')
    p33.style = 'List Number 2'
    
    p34 = doc.add_paragraph('1.2 技术要求')
    p34.style = 'List Number 2'
    
    p35 = doc.add_paragraph('第二条 付款方式')
    p35.style = 'List Number'
    
    p36 = doc.add_paragraph('2.1 预付款')
    p36.style = 'List Number 2'
    
    p37 = doc.add_paragraph('2.2 尾款')
    p37.style = 'List Number 2'
    
    # 保存文档
    doc.save('test-word-advanced-numbering.docx')
    print("高级Word测试文档已创建：test-word-advanced-numbering.docx")

if __name__ == "__main__":
    create_advanced_word_document()
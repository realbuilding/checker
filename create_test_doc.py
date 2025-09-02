from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_document_with_numbering():
    """创建包含Word自动编号的测试文档"""
    doc = Document()
    
    # 添加标题
    doc.add_heading('Word自动编号测试文档', 0)
    
    # 添加说明
    doc.add_paragraph('这是一个测试Word自动编号识别的文档，包含各种常见的编号格式：')
    
    # 一级标题 - 使用Word的自动编号
    doc.add_heading('一级标题示例', level=1)
    
    # 创建带编号的段落
    p1 = doc.add_paragraph('这是第一条内容，使用Word自动编号')
    p1.style = 'List Number'
    
    p2 = doc.add_paragraph('这是第二条内容，使用Word自动编号')
    p2.style = 'List Number'
    
    p3 = doc.add_paragraph('这是第三条内容，使用Word自动编号')
    p3.style = 'List Number'
    
    # 二级标题
    doc.add_heading('二级标题示例', level=2)
    
    # 创建二级编号
    p4 = doc.add_paragraph('这是二级编号的第一项')
    p4.style = 'List Number 2'
    
    p5 = doc.add_paragraph('这是二级编号的第二项')
    p5.style = 'List Number 2'
    
    # 三级标题
    doc.add_heading('三级标题示例', level=3)
    
    # 创建三级编号
    p6 = doc.add_paragraph('这是三级编号的第一项')
    p6.style = 'List Number 3'
    
    p7 = doc.add_paragraph('这是三级编号的第二项')
    p7.style = 'List Number 3'
    
    # 混合编号示例
    doc.add_heading('混合编号示例', level=1)
    
    # 重新开始一级编号
    p8 = doc.add_paragraph('重新开始的一级编号')
    p8.style = 'List Number'
    
    # 添加子项目
    p9 = doc.add_paragraph('子项目A')
    p9.style = 'List Number 2'
    
    p10 = doc.add_paragraph('子项目B')
    p10.style = 'List Number 2'
    
    # 继续一级编号
    p11 = doc.add_paragraph('继续一级编号的第二项')
    p11.style = 'List Number'
    
    # 添加一些特殊格式的内容
    doc.add_heading('特殊格式示例', level=1)
    
    # 带编号的标题段落
    p12 = doc.add_paragraph('1. 带数字编号的段落')
    p13 = doc.add_paragraph('2. 带数字编号的段落')
    p14 = doc.add_paragraph('3. 带数字编号的段落')
    
    # 带括号的编号
    p15 = doc.add_paragraph('(1) 带括号的编号')
    p16 = doc.add_paragraph('(2) 带括号的编号')
    p17 = doc.add_paragraph('(3) 带括号的编号')
    
    # 字母编号
    p18 = doc.add_paragraph('A. 字母编号')
    p19 = doc.add_paragraph('B. 字母编号')
    p20 = doc.add_paragraph('C. 字母编号')
    
    # 保存文档
    doc.save('test-word-numbering.docx')
    print("Word文档已创建：test-word-numbering.docx")

if __name__ == "__main__":
    create_word_document_with_numbering()